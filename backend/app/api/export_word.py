"""数据导出 API — 导出为 Word 文档。"""
from datetime import date, datetime, timedelta, timezone
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_db
from app.models import (
    Project,
    Experience,
    Issue,
    Solution,
    Knowledge,
    Decision,
    Review,
)

router = APIRouter(prefix="/export", tags=["export"])

# 北京时间 UTC+8
BEIJING_TZ = timezone(timedelta(hours=8))


def _now_beijing() -> datetime:
    """获取当前北京时间。"""
    return datetime.now(BEIJING_TZ)


def _serialize_row(row: Any) -> dict:
    """将 SQLAlchemy row 对象转为 dict。"""
    d = {}
    for col in row.__table__.columns:
        val = getattr(row, col.name)
        if isinstance(val, datetime):
            val = val.strftime("%Y-%m-%d %H:%M")
        elif isinstance(val, date):
            val = val.strftime("%Y-%m-%d")
        d[col.name] = val
    return d


def _setup_styles(doc: Document):
    """设置文档样式，统一所有标题字体和大小。"""
    FONT_NAME = 'Microsoft YaHei'

    # 正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)

    # 统一所有 Heading 样式的字体和大小
    heading_sizes = {0: 22, 1: 16, 2: 13, 3: 11}
    for level, size in heading_sizes.items():
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = FONT_NAME
            heading_style.font.size = Pt(size)
        except KeyError:
            pass


def _add_section_header(doc: Document, title: str, color: tuple = (26, 86, 219)):
    """添加带颜色的章节标题。"""
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """添加格式化表格。"""
    if not rows:
        doc.add_paragraph('暂无数据', style='Intense Quote')
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val) if val else '—'
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # 表格后空行


def _build_full_doc(data: dict) -> Document:
    """构建完整导出文档。"""
    doc = Document()
    _setup_styles(doc)

    # 标题
    title = doc.add_heading('CaptureOS 数据导出报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 导出信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'导出时间：{_now_beijing().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # === 项目 ===
    _add_section_header(doc, '📁 项目')
    _add_table(doc,
        ['标题', '状态', '优先级', '技术栈', '描述'],
        [[p['title'], p['status'], p['priority'], p.get('tech_stack', ''), p.get('description', '') or ''] for p in data.get('projects', [])]
    )

    # === 经验 ===
    _add_section_header(doc, '💡 经验')
    for exp in data.get('experiences', []):
        doc.add_heading(exp['title'], level=2)
        if exp.get('summary'):
            doc.add_paragraph(f'摘要：{exp["summary"]}')
        if exp.get('context'):
            doc.add_paragraph(f'背景：{exp["context"]}')
        if exp.get('result'):
            doc.add_paragraph(f'结果：{exp["result"]}')
        if exp.get('lesson'):
            doc.add_paragraph(f'经验教训：{exp["lesson"]}')
        doc.add_paragraph()

    # === 问题 ===
    _add_section_header(doc, '⚠️ 问题')
    _add_table(doc,
        ['标题', '状态', '优先级', '根因', '发现日期'],
        [[i['title'], i['status'], i['priority'], i.get('root_cause', '') or '—',
          i.get('discovered_date', '') or '—'] for i in data.get('issues', [])]
    )

    # === 解决方案 ===
    _add_section_header(doc, '🔧 解决方案')
    for sol in data.get('solutions', []):
        doc.add_heading(sol['title'], level=2)
        if sol.get('approach'):
            doc.add_paragraph(f'方案：{sol["approach"]}')
        if sol.get('outcome'):
            doc.add_paragraph(f'结果：{sol["outcome"]}')
        if sol.get('effectiveness'):
            doc.add_paragraph(f'有效性：{"⭐" * sol["effectiveness"]}')
        doc.add_paragraph()

    # === 知识 ===
    _add_section_header(doc, '📚 知识')
    for k in data.get('knowledge', []):
        doc.add_heading(k['title'], level=2)
        meta = doc.add_paragraph()
        meta.add_run(f'分类：{k.get("category", "—")}　|　可信度：{k.get("confidence", 0)}/10　|　状态：{k.get("status", "—")}')
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        if k.get('content'):
            content = doc.add_paragraph(k['content'])
            content.paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph()

    # === 决策 ===
    _add_section_header(doc, '🎯 决策')
    for d in data.get('decisions', []):
        doc.add_heading(d['title'], level=2)
        if d.get('background'):
            doc.add_paragraph(f'背景：{d["background"]}')
        if d.get('options'):
            opts = d['options']
            if isinstance(opts, list):
                doc.add_paragraph('备选方案：')
                for opt in opts:
                    doc.add_paragraph(f'• {opt}', style='List Bullet')
        if d.get('reason'):
            doc.add_paragraph(f'原因：{d["reason"]}')
        if d.get('result'):
            doc.add_paragraph(f'结果：{d["result"]}')
        doc.add_paragraph()

    # === 复盘 ===
    _add_section_header(doc, '🔄 复盘')
    for r in data.get('reviews', []):
        doc.add_heading(r['title'], level=2)
        if r.get('rating'):
            doc.add_paragraph(f'评分：{"⭐" * r["rating"]}')
        if r.get('success_factors'):
            factors = r['success_factors']
            if isinstance(factors, list):
                doc.add_paragraph('成功因素：')
                for f in factors:
                    doc.add_paragraph(f'✅ {f}', style='List Bullet')
        if r.get('improvements'):
            imps = r['improvements']
            if isinstance(imps, list):
                doc.add_paragraph('改进方向：')
                for imp in imps:
                    doc.add_paragraph(f'🔧 {imp}', style='List Bullet')
        doc.add_paragraph()

    return doc


def _build_single_doc(entity_type: str, items: list[dict]) -> Document:
    """构建单个模块的导出文档。"""
    doc = Document()
    _setup_styles(doc)

    LABELS = {
        'projects': ('📁 项目', (26, 86, 219)),
        'experiences': ('💡 经验', (16, 185, 129)),
        'issues': ('⚠️ 问题', (245, 158, 11)),
        'solutions': ('🔧 解决方案', (139, 92, 246)),
        'knowledge': ('📚 知识', (59, 130, 246)),
        'decisions': ('🎯 决策', (239, 68, 68)),
        'reviews': ('🔄 复盘', (16, 185, 129)),
    }

    label, color = LABELS.get(entity_type, (entity_type, (26, 86, 219)))

    title = doc.add_heading(f'{label}数据导出', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'导出时间：{_now_beijing().strftime("%Y-%m-%d %H:%M")}　|　共 {len(items)} 条')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    if entity_type == 'projects':
        _add_table(doc,
            ['标题', '状态', '优先级', '技术栈', '描述'],
            [[p['title'], p['status'], p['priority'], p.get('tech_stack', ''), p.get('description', '') or ''] for p in items]
        )
    elif entity_type == 'experiences':
        for item in items:
            doc.add_heading(item['title'], level=2)
            for field, label in [('summary', '摘要'), ('context', '背景'), ('result', '结果'), ('lesson', '经验教训')]:
                if item.get(field):
                    doc.add_paragraph(f'{label}：{item[field]}')
            doc.add_paragraph()
    elif entity_type == 'issues':
        _add_table(doc,
            ['标题', '状态', '优先级', '根因', '发现日期'],
            [[i['title'], i['status'], i['priority'], i.get('root_cause', '') or '—',
              i.get('discovered_date', '') or '—'] for i in items]
        )
    elif entity_type == 'solutions':
        for item in items:
            doc.add_heading(item['title'], level=2)
            for field, label in [('approach', '方案'), ('outcome', '结果')]:
                if item.get(field):
                    doc.add_paragraph(f'{label}：{item[field]}')
            if item.get('effectiveness'):
                doc.add_paragraph(f'有效性：{"⭐" * item["effectiveness"]}')
            doc.add_paragraph()
    elif entity_type == 'knowledge':
        for item in items:
            doc.add_heading(item['title'], level=2)
            meta = doc.add_paragraph()
            meta.add_run(f'分类：{item.get("category", "—")}　|　可信度：{item.get("confidence", 0)}/10　|　状态：{item.get("status", "—")}')
            meta.runs[0].font.size = Pt(9)
            meta.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            if item.get('content'):
                content = doc.add_paragraph(item['content'])
                content.paragraph_format.left_indent = Inches(0.3)
            doc.add_paragraph()
    elif entity_type == 'decisions':
        for item in items:
            doc.add_heading(item['title'], level=2)
            for field, label in [('background', '背景'), ('reason', '原因'), ('result', '结果')]:
                if item.get(field):
                    doc.add_paragraph(f'{label}：{item[field]}')
            doc.add_paragraph()
    elif entity_type == 'reviews':
        for item in items:
            doc.add_heading(item['title'], level=2)
            if item.get('rating'):
                doc.add_paragraph(f'评分：{"⭐" * item["rating"]}')
            for field, label in [('success_factors', '成功因素'), ('improvements', '改进方向')]:
                vals = item.get(field)
                if vals and isinstance(vals, list):
                    doc.add_paragraph(f'{label}：')
                    for v in vals:
                        doc.add_paragraph(f'• {v}', style='List Bullet')
            doc.add_paragraph()

    return doc


async def _fetch_all(db: AsyncSession, model) -> list[dict]:
    result = await db.execute(select(model))
    return [_serialize_row(r) for r in result.scalars().all()]


@router.get("/word/")
async def export_all_word(types: str = "", db: AsyncSession = Depends(get_db)):
    """导出数据为 Word 文档。types 参数可选，逗号分隔的实体类型列表。"""
    ALL_TYPES = ["projects", "experiences", "issues", "solutions", "knowledge", "decisions", "reviews"]
    MODEL_MAP_LOCAL = {
        "projects": Project,
        "experiences": Experience,
        "issues": Issue,
        "solutions": Solution,
        "knowledge": Knowledge,
        "decisions": Decision,
        "reviews": Review,
    }

    if types:
        selected = [t.strip() for t in types.split(",") if t.strip() in ALL_TYPES]
    else:
        selected = ALL_TYPES

    data = {}
    for t in selected:
        data[t] = await _fetch_all(db, MODEL_MAP_LOCAL[t])

    doc = _build_full_doc(data)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    today = _now_beijing().strftime("%Y-%m-%d")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=captureos-export-{today}.docx"},
    )


MODEL_MAP = {
    "projects": Project,
    "experiences": Experience,
    "issues": Issue,
    "solutions": Solution,
    "knowledge": Knowledge,
    "decisions": Decision,
    "reviews": Review,
}


@router.get("/word/{entity_type}")
async def export_entity_word(entity_type: str, db: AsyncSession = Depends(get_db)):
    """导出单个模块全部数据为 Word 文档。"""
    model = MODEL_MAP.get(entity_type)
    if not model:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=f"未知实体类型: {entity_type}")

    items = await _fetch_all(db, model)
    doc = _build_single_doc(entity_type, items)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    today = _now_beijing().strftime("%Y-%m-%d")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=captureos-{entity_type}-{today}.docx"},
    )


@router.post("/word/{entity_type}/selected")
async def export_selected_word(
    entity_type: str,
    payload: dict,
    db: AsyncSession = Depends(get_db),
):
    """导出选中的记录为 Word 文档。"""
    from fastapi import HTTPException
    from sqlalchemy import select as sel

    model = MODEL_MAP.get(entity_type)
    if not model:
        raise HTTPException(status_code=400, detail=f"未知实体类型: {entity_type}")

    ids = payload.get("ids", [])
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要导出的记录")

    result = await db.execute(sel(model).where(model.id.in_(ids)))
    items = [_serialize_row(r) for r in result.scalars().all()]

    if not items:
        raise HTTPException(status_code=404, detail="未找到选中的记录")

    doc = _build_single_doc(entity_type, items)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    today = _now_beijing().strftime("%Y-%m-%d")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=captureos-{entity_type}-selected-{today}.docx"},
    )
